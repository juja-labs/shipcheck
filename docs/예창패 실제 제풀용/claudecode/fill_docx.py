#!/usr/bin/env python3
"""
예비창업패키지 사업계획서 양식을 draft 내용으로 채우는 스크립트
"""

import re
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = Path(__file__).parent
TEMPLATE = BASE_DIR / '2026년 예비창업패키지 사업계획서 양식.docx'
DRAFT = BASE_DIR / '예창패_초안_draft.md'
OUTPUT = BASE_DIR / '2026년_예비창업패키지_사업계획서_Personica.docx'

FONT_NAME = '맑은 고딕'
BODY_FONT_SIZE = 10  # pt
SMALL_FONT_SIZE = 9

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_run_font(run, size=BODY_FONT_SIZE, bold=False, font_name=FONT_NAME):
    """Set font properties on a run."""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_cell_text(cell, text, bold=False, size=BODY_FONT_SIZE):
    """Set cell text, preserving cell formatting."""
    # Clear all paragraphs except first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    para = cell.paragraphs[0]
    # Clear existing runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    # Also clear any remaining text nodes
    for child in list(para._element):
        if child.tag == qn('w:r'):
            para._element.remove(child)

    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)


def create_para(doc_body, text='', bold=False, size=BODY_FONT_SIZE, indent_cm=0,
                spacing_before=0, spacing_after=0, alignment=None):
    """Create a new paragraph element with formatting."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # Spacing
    if spacing_before or spacing_after:
        spacing = OxmlElement('w:spacing')
        if spacing_before:
            spacing.set(qn('w:before'), str(int(spacing_before * 20)))  # twips
        if spacing_after:
            spacing.set(qn('w:after'), str(int(spacing_after * 20)))
        spacing.set(qn('w:line'), '276')  # 1.15 line spacing
        pPr.append(spacing)
    else:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '40')
        spacing.set(qn('w:line'), '276')
        pPr.append(spacing)

    # Indentation
    if indent_cm > 0:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(indent_cm * 567)))  # cm to twips
        pPr.append(ind)

    # Alignment
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)

    p.append(pPr)

    if text:
        add_formatted_text(p, text, bold=bold, size=size)

    return p


def add_formatted_text(p_element, text, bold=False, size=BODY_FONT_SIZE):
    """Add text with inline **bold** parsing to a paragraph element."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _add_run(p_element, part[2:-2], bold=True, size=size)
        else:
            _add_run(p_element, part, bold=bold, size=size)


def _add_run(p_element, text, bold=False, size=BODY_FONT_SIZE):
    """Add a run element to a paragraph."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_element.append(r)


def insert_after(ref_element, new_element):
    """Insert new_element right after ref_element in the XML tree."""
    ref_element.addnext(new_element)
    return new_element


def remove_empty_paragraphs(body, start_elem, end_elem):
    """Remove empty paragraphs between start (exclusive) and end (exclusive)."""
    elements = list(body)
    start_idx = elements.index(start_elem)
    end_idx = elements.index(end_elem) if end_elem is not None else len(elements)

    to_remove = []
    for i in range(start_idx + 1, end_idx):
        elem = elements[i]
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            all_text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
            if not all_text.strip():
                to_remove.append(elem)

    for elem in to_remove:
        body.remove(elem)


def create_docx_table(doc, rows_data, col_widths=None, header_row=True):
    """Create a docx table and return its element.
    rows_data: list of lists of strings (cell text)
    """
    n_rows = len(rows_data)
    n_cols = max(len(r) for r in rows_data) if rows_data else 1

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= n_cols:
                break
            cell = table.cell(ri, ci)
            # Clear default paragraph
            para = cell.paragraphs[0]
            for run in list(para.runs):
                run._element.getparent().remove(run._element)

            # Parse bold
            is_header = header_row and ri == 0
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    set_run_font(run, size=SMALL_FONT_SIZE, bold=True)
                else:
                    run = para.add_run(part)
                    set_run_font(run, size=SMALL_FONT_SIZE, bold=is_header)

            # Set spacing
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '240')
            pPr.append(spacing)

    # Set column widths if provided
    if col_widths:
        for ri in range(n_rows):
            for ci, width in enumerate(col_widths):
                if ci < n_cols:
                    tc = table.cell(ri, ci)._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(int(width * 567)))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)

    return table


# ============================================================
# MARKDOWN PARSER
# ============================================================

def parse_markdown_to_elements(doc, body, markdown_text, font_size=BODY_FONT_SIZE):
    """Parse markdown text and return list of XML elements to insert."""
    elements = []
    lines = markdown_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip --- dividers
        if stripped == '---':
            i += 1
            continue

        # Section headers ###
        if stripped.startswith('### '):
            title = stripped[4:]
            # Remove markdown bold from title for display
            clean_title = title.replace('**', '')
            p = create_para(body, bold=True, size=font_size + 1, spacing_before=6)
            add_formatted_text(p, title, bold=True, size=font_size + 1)
            elements.append(('p', p))
            i += 1
            continue

        # Section headers ##
        if stripped.startswith('## '):
            title = stripped[3:]
            p = create_para(body, bold=True, size=font_size + 2, spacing_before=8)
            add_formatted_text(p, title, bold=True, size=font_size + 2)
            elements.append(('p', p))
            i += 1
            continue

        # Blockquotes > (시각자료 placeholders etc.)
        if stripped.startswith('> '):
            text = stripped[2:]
            p = create_para(body, size=font_size - 1, indent_cm=0.5)
            add_formatted_text(p, f'[{text}]', size=font_size - 1)
            elements.append(('p', p))
            i += 1
            continue

        # Code blocks ``` ... ```
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```

            # Create paragraph with code
            p = create_para(body, indent_cm=0.3)
            for j, code_line in enumerate(code_lines):
                if j > 0:
                    # Line break
                    r = OxmlElement('w:r')
                    br = OxmlElement('w:br')
                    r.append(br)
                    p.append(r)
                _add_run(p, code_line or ' ', size=font_size - 1)
            elements.append(('p', p))
            continue

        # Tables | ... |
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows = []
            for tl in table_lines:
                # Skip separator rows like |---|---|
                cells_text = tl.strip('|').strip()
                if re.match(r'^[\s\-:|\s]+$', cells_text):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)

            if rows:
                elements.append(('table', rows))
            continue

        # Bullet lists -
        if stripped.startswith('- '):
            text = stripped[2:]
            p = create_para(body, indent_cm=0.5, size=font_size)
            _add_run(p, '• ', size=font_size)
            add_formatted_text(p, text, size=font_size)
            elements.append(('p', p))
            i += 1
            continue

        # Numbered items (like **1) text** or 1. text)
        # Regular paragraph
        p = create_para(body, size=font_size)
        add_formatted_text(p, stripped, size=font_size)
        elements.append(('p', p))
        i += 1

    return elements


def insert_elements_after(doc, body, anchor_elem, elements):
    """Insert parsed elements after anchor element."""
    current = anchor_elem

    for elem_type, elem_data in elements:
        if elem_type == 'p':
            current = insert_after(current, elem_data)
        elif elem_type == 'table':
            rows = elem_data
            n_rows = len(rows)
            n_cols = max(len(r) for r in rows)

            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = 'Table Normal'

            # Add borders to the table
            tbl_elem = table._tbl
            tblPr = tbl_elem.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_elem.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci >= n_cols:
                        break
                    cell = table.cell(ri, ci)
                    para = cell.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)

                    is_header = ri == 0
                    parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            set_run_font(run, size=SMALL_FONT_SIZE, bold=True)
                        else:
                            run = para.add_run(part)
                            set_run_font(run, size=SMALL_FONT_SIZE, bold=is_header)

                    # Compact spacing
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        para._element.insert(0, pPr)
                    sp = OxmlElement('w:spacing')
                    sp.set(qn('w:after'), '0')
                    sp.set(qn('w:line'), '240')
                    pPr.append(sp)

            # Move table to correct position
            tbl_elem = table._tbl
            body.remove(tbl_elem)
            current = insert_after(current, tbl_elem)

    return current


# ============================================================
# READ DRAFT
# ============================================================

def read_draft_sections(draft_path):
    """Read draft markdown and split into sections."""
    with open(draft_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = {}

    # Split by ### headers
    parts = re.split(r'\n(?=### \d+-\d+\.)', content)

    for part in parts:
        match = re.match(r'### (\d+-\d+)\.\s', part)
        if match:
            key = match.group(1)
            sections[key] = part.strip()

    return sections


# ============================================================
# MAIN
# ============================================================

def main():
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    all_elements = list(body)

    # Read draft
    sections = read_draft_sections(DRAFT)

    print("Draft sections found:", list(sections.keys()))

    # --------------------------------------------------------
    # Find key elements by their text content
    # --------------------------------------------------------
    def find_element_by_text(search_text, tag='p'):
        for elem in body:
            elem_tag = elem.tag.split('}')[-1]
            if elem_tag != tag:
                continue
            all_text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
            if search_text in all_text:
                return elem
        return None

    def find_all_elements_by_text(search_text, tag='p'):
        results = []
        for elem in body:
            elem_tag = elem.tag.split('}')[-1]
            if elem_tag != tag:
                continue
            all_text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
            if search_text in all_text:
                results.append(elem)
        return results

    # Find ㅇ markers
    oo_markers = find_all_elements_by_text('ㅇ')
    print(f"Found {len(oo_markers)} ㅇ markers")

    # Find section headers and structural elements
    section1_header = find_element_by_text('1. 문제 인식')
    section2_header = find_element_by_text('2. 실현 가능성')
    section3_header = find_element_by_text('3. 성장전략')
    section4_header = find_element_by_text('4. 팀 구성')
    schedule_header = find_element_by_text('사업추진 일정(협약기간 내)')
    stage1_header = find_element_by_text('1단계 정부지원사업비')
    stage2_header = find_element_by_text('2단계 정부지원사업비')
    full_schedule_header = find_element_by_text('사업추진 일정(전체')
    team_table_header = find_element_by_text('팀 구성(안)')
    partner_header = find_element_by_text('협력 기관 현황')

    # --------------------------------------------------------
    # FILL TABLE 3 (일반현황)
    # --------------------------------------------------------
    print("Filling Table 3...")
    t3 = doc.tables[3]  # 9 rows, 8 cols (일반현황)

    # R0: 창업아이템명
    set_cell_text(t3.rows[0].cells[3],
        '감정·인지·의사결정을 구조적으로 시뮬레이션하는 AI 합성 페르소나 리서치 플랫폼',
        size=BODY_FONT_SIZE)

    # R1: 산출물
    set_cell_text(t3.rows[1].cells[3], '웹사이트(1개)', size=BODY_FONT_SIZE)

    # R2: 직업 + 기업명
    set_cell_text(t3.rows[2].cells[3], '예비창업가', size=BODY_FONT_SIZE)
    set_cell_text(t3.rows[2].cells[6], 'Personica', size=BODY_FONT_SIZE)

    # R5-R8: 팀 구성
    # R5: UX 리서치 전문가
    set_cell_text(t3.rows[5].cells[0], '1', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[5].cells[1], '팀원', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[5].cells[2], 'UX 리서치 전문가', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[5].cells[4], 'UX 리서치 실무 경력 3년 이상', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[5].cells[7], '예정(\'26.06)', size=SMALL_FONT_SIZE)

    # R6: SW 엔지니어
    set_cell_text(t3.rows[6].cells[0], '2', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[6].cells[1], '팀원', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[6].cells[2], 'SW 엔지니어', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[6].cells[4], 'Python/React 개발 경력 3년 이상', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[6].cells[7], '예정(\'26.06)', size=SMALL_FONT_SIZE)

    # R7: 사업화 담당
    set_cell_text(t3.rows[7].cells[0], '3', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[7].cells[1], '팀원', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[7].cells[2], '사업화 및 시장검증', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[7].cells[4], '스타트업 마케팅/그로스 경력', size=SMALL_FONT_SIZE)
    set_cell_text(t3.rows[7].cells[7], '예정(\'26.09)', size=SMALL_FONT_SIZE)

    # R8: clear
    set_cell_text(t3.rows[8].cells[0], '', size=SMALL_FONT_SIZE)

    # --------------------------------------------------------
    # FILL TABLE 4 (개요/요약)
    # --------------------------------------------------------
    print("Filling Table 4...")
    t4 = doc.tables[4]  # 8 rows, 5 cols

    # R0: 명칭 / 범주 headers - already filled, add content
    set_cell_text(t4.rows[0].cells[1], 'Personica', bold=True, size=BODY_FONT_SIZE)
    set_cell_text(t4.rows[0].cells[4], 'AI / 합성 페르소나 시뮬레이션', size=BODY_FONT_SIZE)

    # R1: 아이템 개요
    set_cell_text(t4.rows[1].cells[1],
        'Personica는 감정·인지·의사결정을 구조적으로 시뮬레이션하는 AI 합성 페르소나 플랫폼이다. '
        'AI 페르소나가 서베이에 응답하고, UI/UX를 리뷰하고, 실제 제품을 사용하며 '
        '"사람처럼" 피드백을 생성한다. 기존 리서치(설문·분석·테스트)의 비용·시간·규모 한계를 해결한다.',
        size=SMALL_FONT_SIZE)

    # R2: 문제 인식
    set_cell_text(t4.rows[2].cells[1],
        '스타트업 실패의 42%는 시장 수요 오판(CB Insights). 설문은 의견만, 분석도구는 출시 후만, '
        '사용성 테스트는 비싸고 느림. AI 페르소나가 대안으로 등장(Aaru $1B, Simile $100M 투자)했으나 '
        '긍정 편향·제품 미사용·분화 부족 3가지 한계 존재.',
        size=SMALL_FONT_SIZE)

    # R3: 실현 가능성
    set_cell_text(t4.rows[3].cells[1],
        '페르소나 리얼리티 엔진(5레이어: 성격→인지→감정→판단→기억)으로 3가지 공백 해결. '
        '핵심 엔진 구현 완료, 30명 시뮬레이션 분화 검증(세그먼트별 평점 3.50~4.33), '
        'G2 실제 리뷰 대비 긍정 90%·불만 71% 테마 일치 확인. 감정 엔진 ON/OFF 비교로 긍정 편향 구조적 감소 입증.',
        size=SMALL_FONT_SIZE)

    # R4: 성장전략
    set_cell_text(t4.rows[4].cells[1],
        '건별 과금(PAYG) $19~$499. 합성 서베이·UI/UX 리뷰·제품 체험 3모드. '
        'Growth 30명 기준 마진 75~87%. 한국 PLG 검증(6개월) → Product Hunt 글로벌 확장. '
        '1년차 $240K → 3년차 $7M → 5년차 $24M 매출 목표.',
        size=SMALL_FONT_SIZE)

    # R5: 팀 구성
    set_cell_text(t4.rows[5].cells[1],
        '대표(유성): AI Engineer 11년 8개월(KT 20명 팀 리딩·SKT 생성형AI·SSG 빅데이터). '
        'UX 리서치 전문가 + SW 엔지니어(1개월차) + 사업화 담당(4개월차). '
        '대표가 AI 엔진 직접 구현, UX 전문가가 아웃풋 품질 검증.',
        size=SMALL_FONT_SIZE)

    # R6-R7: 이미지 - leave as placeholder
    set_cell_text(t4.rows[6].cells[1], '[시각자료: 서비스 플로우 다이어그램]', size=SMALL_FONT_SIZE)
    set_cell_text(t4.rows[6].cells[3], '[시각자료: 경쟁사 포지셔닝 맵]', size=SMALL_FONT_SIZE)
    set_cell_text(t4.rows[7].cells[1], '[시각자료: 페르소나 리얼리티 엔진 구조]', size=SMALL_FONT_SIZE)
    set_cell_text(t4.rows[7].cells[3], '[시각자료: 리포트 샘플 목업]', size=SMALL_FONT_SIZE)

    # --------------------------------------------------------
    # FILL STRUCTURAL TABLES 5-10 FIRST (before body insertion changes indices)
    # --------------------------------------------------------

    # Save references to tables by index BEFORE body content insertion
    t5 = doc.tables[5]  # 사업추진 일정
    t6 = doc.tables[6]  # 1단계 집행
    t7 = doc.tables[7]  # 2단계 집행
    t8 = doc.tables[8]  # 전체 일정
    t9 = doc.tables[9]  # 팀 구성
    t10 = doc.tables[10]  # 협력기관

    # --------------------------------------------------------
    # FILL TABLE 5 (사업추진 일정 - 협약기간 내)
    # --------------------------------------------------------
    print("Filling Table 5...")

    schedule_data = [
        ('1', '핵심 엔진 제품화 + 3개 모드 구현', '26.06 ~ 26.08', '페르소나 리얼리티 엔진 고도화, 합성 서베이·UI/UX 리뷰·제품 체험 3모드 개발, 리포트 자동화'),
        ('2', '다중 세션 기억 + 베타 출시', '26.09 ~ 26.10', '다중 세션 기억·학습 기능, 리포트 고도화, 베타 버전 공개'),
        ('3', '정식 출시 + 글로벌 PLG', '26.11 ~ 27.01', '셀프서브 온보딩, 결제 연동, Product Hunt·Show HN 출시'),
        ('4', '초기 고객 확보 + 유료 전환', '26.06 ~ 27.01', '고객 인터뷰 15건, 무료 체험 50건+, 유료 전환 10~20건'),
    ]

    for i, (num, content, period, detail) in enumerate(schedule_data):
        row = t5.rows[i + 1]
        set_cell_text(row.cells[0], num, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[1], content, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[2], period, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[3], detail, size=SMALL_FONT_SIZE)

    # Clear last example row
    if len(t5.rows) > 5:
        for cell in t5.rows[5].cells:
            set_cell_text(cell, '', size=SMALL_FONT_SIZE)

    # --------------------------------------------------------
    # FILL TABLE 6 (1단계 집행계획)
    # --------------------------------------------------------
    print("Filling Table 6...")

    # 1단계: 1~5개월차 (2026.06~10) ≈ 2,000만원
    stage1_data = [
        ('인건비', '▪ 대표 AI 엔진 개발 (5개월×100만)', '5,000,000'),
        ('인건비', '▪ UX 리서치 전문가 (3개월×100만+2개월×80만)', '4,600,000'),
        ('인건비', '▪ SW 엔지니어 (3개월×130만+2개월×150만)', '6,900,000'),
        ('인건비', '▪ 사업화 담당 (2개월×100만, 4~5개월차)', '2,000,000'),
        ('지급수수료', '▪ 클라우드·LLM API 비용 (5개월)', '2,000,000'),
        ('지급수수료', '▪ 협업도구·법무·세무 (5개월)', '1,000,000'),
        ('여비', '▪ 고객 인터뷰·미팅 교통비 (15건)', '1,200,000'),
    ]

    # Adjust table rows
    while len(t6.rows) < len(stage1_data) + 2:  # +1 header +1 total
        t6.add_row()

    for i, (category, basis, amount) in enumerate(stage1_data):
        row = t6.rows[i + 1]
        amt = int(amount.replace(',', ''))
        set_cell_text(row.cells[0], category, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[1], basis, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[2], f'{amt:,}', size=SMALL_FONT_SIZE)

    # Total row
    total_row = t6.rows[len(stage1_data) + 1]
    set_cell_text(total_row.cells[0], '합 계', bold=True, size=SMALL_FONT_SIZE)
    set_cell_text(total_row.cells[1], '', size=SMALL_FONT_SIZE)
    total1 = sum(int(d[2].replace(',', '')) for d in stage1_data)
    set_cell_text(total_row.cells[2], f'{total1:,}', bold=True, size=SMALL_FONT_SIZE)

    # Remove extra rows
    while len(t6.rows) > len(stage1_data) + 2:
        tr = t6.rows[-1]._tr
        t6._tbl.remove(tr)

    # --------------------------------------------------------
    # FILL TABLE 7 (2단계 집행계획)
    # --------------------------------------------------------
    print("Filling Table 7...")

    # 2단계: 6~8개월차 (2026.11~2027.01) ≈ 2,000만원 (잔액)
    remaining = 40000000 - total1

    stage2_data = [
        ('인건비', '▪ 대표 AI 엔진 개발 (3개월×100만)', '3,000,000'),
        ('인건비', '▪ UX 리서치 전문가 (3개월×80만)', '2,400,000'),
        ('인건비', '▪ SW 엔지니어 (3개월×150만)', '4,500,000'),
        ('인건비', '▪ 사업화 담당 (3개월×110만)', '3,300,000'),
        ('지급수수료', '▪ 클라우드·LLM API 비용 (3개월)', '1,300,000'),
        ('지급수수료', '▪ 협업도구·법무·세무 (3개월)', '600,000'),
        ('광고선전비', '▪ 초기 마케팅 실험·콘텐츠 제작', '1,600,000'),
        ('여비', '▪ 고객 미팅·전시회 교통비', '600,000'),
    ]

    while len(t7.rows) < len(stage2_data) + 2:
        t7.add_row()

    for i, (category, basis, amount) in enumerate(stage2_data):
        row = t7.rows[i + 1]
        amt = int(amount.replace(',', ''))
        set_cell_text(row.cells[0], category, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[1], basis, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[2], f'{amt:,}', size=SMALL_FONT_SIZE)

    total_row = t7.rows[len(stage2_data) + 1]
    set_cell_text(total_row.cells[0], '합 계', bold=True, size=SMALL_FONT_SIZE)
    set_cell_text(total_row.cells[1], '', size=SMALL_FONT_SIZE)
    total2 = sum(int(d[2].replace(',', '')) for d in stage2_data)
    set_cell_text(total_row.cells[2], f'{total2:,}', bold=True, size=SMALL_FONT_SIZE)

    while len(t7.rows) > len(stage2_data) + 2:
        tr = t7.rows[-1]._tr
        t7._tbl.remove(tr)

    print(f"  1단계: {total1:,}원, 2단계: {total2:,}원, 합계: {total1+total2:,}원")

    # --------------------------------------------------------
    # FILL TABLE 8 (전체 사업단계 일정)
    # --------------------------------------------------------
    print("Filling Table 8...")

    full_schedule = [
        ('1', '핵심 엔진 제품화 (알파)', '26년 06~08월', '3개 모드 개발, 리포트 자동화, 웹 서비스 구축'),
        ('2', '베타 출시 + 한국 PLG', '26년 09~10월', '다중 세션 기억, 유료 전환 시작, 커뮤니티 PLG'),
        ('3', '정식 출시 + 글로벌 확장', '26년 11~27년 01월', 'Product Hunt 출시, 유료 10~20건'),
        ('4', '글로벌 PLG 성장', '27년 상반기', '글로벌 매출 70~80%, 시드 투자 추진'),
        ('5', 'SOM 진입 + 모드 확장', '28~29년', 'A/B 테스트·광고 리서치 모드, $7M 매출'),
    ]

    while len(t8.rows) < len(full_schedule) + 1:
        t8.add_row()

    for i, (num, content, period, detail) in enumerate(full_schedule):
        row = t8.rows[i + 1]
        set_cell_text(row.cells[0], num, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[1], content, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[2], period, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[3], detail, size=SMALL_FONT_SIZE)

    while len(t8.rows) > len(full_schedule) + 1:
        tr = t8.rows[-1]._tr
        t8._tbl.remove(tr)

    # --------------------------------------------------------
    # FILL TABLE 9 (팀 구성안)
    # --------------------------------------------------------
    print("Filling Table 9...")

    team_data = [
        ('1', '팀원', 'UX 리서치 전문가\n(시뮬레이션 아웃풋 품질 검증)', 'UX 리서치 실무 경력 3년 이상, 리서치 설계·분석 전문', '예정(\'26.06)'),
        ('2', '팀원', 'SW 엔지니어\n(프론트엔드·서비스 개발)', 'Python/React/FastAPI 개발 경력 3년 이상', '예정(\'26.06)'),
        ('3', '팀원', '사업화 및 시장검증\n(마케팅·전환·콘텐츠)', '스타트업 마케팅/그로스 경력, PLG 운영 경험', '예정(\'26.09)'),
    ]

    while len(t9.rows) < len(team_data) + 1:
        t9.add_row()

    for i, (num, position, role, capability, status) in enumerate(team_data):
        row = t9.rows[i + 1]
        set_cell_text(row.cells[0], num, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[1], position, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[2], role, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[3], capability, size=SMALL_FONT_SIZE)
        set_cell_text(row.cells[4], status, size=SMALL_FONT_SIZE)

    while len(t9.rows) > len(team_data) + 1:
        tr = t9.rows[-1]._tr
        t9._tbl.remove(tr)

    # --------------------------------------------------------
    # FILL TABLE 10 (협력 기관) - 비움
    # --------------------------------------------------------
    print("Filling Table 10...")

    # Clear example data
    for ri in range(1, len(t10.rows)):
        for cell in t10.rows[ri].cells:
            set_cell_text(cell, '—', size=SMALL_FONT_SIZE)
    set_cell_text(t10.rows[1].cells[0], '—', size=SMALL_FONT_SIZE)

    # --------------------------------------------------------
    # BODY SECTIONS - Insert content at ㅇ markers
    # (Done AFTER table filling to avoid index shifting)
    # --------------------------------------------------------

    section_mapping = {
        0: ['1-1'],
        1: ['1-2'],
        2: ['1-3', '1-4'],
        3: ['2-1', '2-2', '2-3', '2-4', '2-5'],
        4: ['3-1'],
        5: ['3-2', '3-3'],
        6: ['4-1', '4-2', '4-3'],
    }

    # Determine end elements for each ㅇ marker (for removing empty paragraphs)
    end_elements = {}
    if len(oo_markers) >= 7:
        end_elements[0] = oo_markers[1]
        end_elements[1] = oo_markers[2]
        end_elements[2] = section2_header
        end_elements[3] = schedule_header
        end_elements[4] = oo_markers[5]
        end_elements[5] = full_schedule_header
        end_elements[6] = team_table_header

    # Process from bottom to top to maintain element positions
    for marker_idx in sorted(section_mapping.keys(), reverse=True):
        section_keys = section_mapping[marker_idx]
        marker_elem = oo_markers[marker_idx]
        end_elem = end_elements.get(marker_idx)

        print(f"Processing ㅇ marker {marker_idx} → sections {section_keys}")

        combined_md = []
        for key in section_keys:
            if key in sections:
                combined_md.append(sections[key])
            else:
                print(f"  WARNING: Section {key} not found in draft!")

        full_md = '\n\n'.join(combined_md)

        if end_elem is not None:
            remove_empty_paragraphs(body, marker_elem, end_elem)

        for t_elem in marker_elem.iter(qn('w:t')):
            t_elem.text = ''

        parsed = parse_markdown_to_elements(doc, body, full_md)
        insert_elements_after(doc, body, marker_elem, parsed)

    # Remove instruction paragraphs (※ ...)
    instruction_paragraphs = find_all_elements_by_text('※ 개발하고자 하는')
    instruction_paragraphs += find_all_elements_by_text('※ 아이디어를 제품')
    instruction_paragraphs += find_all_elements_by_text('※ 경쟁제품·경쟁사')
    instruction_paragraphs += find_all_elements_by_text('※ 대표자 보유 역량')
    for p in instruction_paragraphs:
        body.remove(p)

    # --------------------------------------------------------
    # SAVE
    # --------------------------------------------------------
    print(f"\nSaving to {OUTPUT}...")
    doc.save(str(OUTPUT))
    print("Done!")
    print(f"\n총 예산: {total1+total2:,}원 (1단계 {total1:,} + 2단계 {total2:,})")


if __name__ == '__main__':
    main()
