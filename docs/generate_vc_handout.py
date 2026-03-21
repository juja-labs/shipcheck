#!/usr/bin/env python3
"""ShipCheck VC 미팅 핸드아웃 Word 문서 생성."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from pathlib import Path

MOCKUPS = Path(__file__).parent / "mockups"
OUTPUT = Path(__file__).parent / "ShipCheck_VC_Handout.docx"


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_body(doc, text, bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_image_full(doc, path, caption=None, width=Inches(6.5)):
    """이미지 + 캡션."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.italic = True


def build():
    doc = Document()

    # 페이지 설정: A4 가로
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 기본 폰트
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    # ─── 표지 ───
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ShipCheck")
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC8, 0x6B, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("AI 페르소나가 제품을 직접 사용하고 리뷰하는 SaaS")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('"만드는 건 끝났어. 이거 사람들이 원할까?"')
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026.03.18  |  Product Overview")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ─── 1. 문제 정의 ───
    add_heading_styled(doc, "1. Problem — 출시 전 사용자 검증의 공백", 1)

    add_body(doc,
        "제품 팀은 페르소나와 유저 저니맵을 기반으로 설계하지만, "
        "실제 다양한 사용자가 쓸 때는 예상과 다르게 동작합니다. "
        "출시 전에 이를 검증할 현실적인 방법이 없습니다.", size=12)

    # 표: 기존 방법 vs ShipCheck
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Shading Accent 1"

    headers = ["", "기존 방법", "ShipCheck"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    rows_data = [
        ["비용", "사용자 리서치 1회 $5,000~30,000", "페르소나 20명 시뮬레이션 ~$15"],
        ["시간", "리크루팅+인터뷰 2~4주", "30분 내 결과"],
        ["실제 사용", "설문/인터뷰 (사용 안 함)", "Playwright로 제품 직접 조작"],
    ]
    for r_idx, row_data in enumerate(rows_data, 1):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = val

    doc.add_paragraph()
    doc.add_page_break()

    # ─── 2. 핵심 기술 — 라이브 대시보드 ───
    add_heading_styled(doc, "2. 핵심 기술 — 5개 AI 페르소나 동시 실행", 1)

    add_body(doc,
        "각 페르소나는 독립된 Claude Code CLI 인스턴스로 실행됩니다. "
        "Playwright로 실제 제품을 조작하면서, 결정론적 감정 엔진(OCC → PAD → SDE)이 "
        "매 행동마다 감정 상태를 수학적으로 계산합니다. "
        "아래는 5개 페르소나가 tally.so를 동시에 사용하는 라이브 대시보드입니다.", size=11)

    add_image_full(doc,
        MOCKUPS / "image (1).png",
        "ShipCheck Live — 5개 페르소나 병렬 실행. 같은 제품이지만 감정 궤적이 전부 다릅니다.",
        width=Inches(9.5))

    doc.add_page_break()

    # ─── 3. 시뮬레이션 뷰 ───
    add_heading_styled(doc, "3. 시뮬레이션 뷰 — 페르소나의 체험을 추적", 1)

    add_body(doc,
        "각 페르소나의 행동을 타임라인으로 추적합니다. "
        "왼쪽은 Action Trace (클릭, 스크롤, 입력 등), 가운데는 실제 화면, "
        "오른쪽은 페르소나의 머릿속 (감정, 혼란 수준, 하고 싶은 것)입니다.", size=11)

    add_image_full(doc,
        MOCKUPS / "v4-redesign-sim.png",
        "박서연 (52세, 초등학교 교사, 불안형) — 템플릿 페이지에서 좌절. '한국어가 하나도 없어서 뭘 눌러야 할지 모르겠다'",
        width=Inches(9.5))

    doc.add_page_break()

    # ─── 4. 세션 리플레이 ───
    add_heading_styled(doc, "4. 세션 리플레이 — 영상처럼 재생", 1)

    add_body(doc,
        "페르소나의 전체 세션을 영상처럼 재생할 수 있습니다. "
        "각 스텝마다 실제 화면 + 페르소나의 감정 말풍선 + PAD 감정 바가 동기화됩니다. "
        "PM이 '이 사용자가 왜 이탈했는지'를 직접 체감할 수 있습니다.", size=11)

    add_image_full(doc,
        MOCKUPS / "v4-redesign-replay.png",
        "세션 리플레이 — 박서연이 탈락하는 순간. PAD -0.6 (좌절). '가입하기는 싫고... 그냥 구글 폼으로 돌아가자.'",
        width=Inches(9.5))

    doc.add_page_break()

    # ─── 5. 경쟁사 비교 ───
    add_heading_styled(doc, "5. 경쟁사 비교 — A/B 시뮬레이션", 1)

    add_body(doc,
        "같은 페르소나 세트로 고객 제품과 경쟁사를 모두 시뮬레이션합니다. "
        "동일 조건에서의 핵심 플로우 완료율, 만족도, NPS, 이탈률을 비교합니다. "
        "실제 사용자 테스트 없이 경쟁 벤치마크를 얻을 수 있습니다.", size=11)

    add_image_full(doc,
        MOCKUPS / "v4-redesign-bench.png",
        "tally.so vs Typeform — 동일 19명 페르소나 시뮬레이션. tally.so 완료율 63% vs Typeform 89%",
        width=Inches(9.5))

    doc.add_page_break()

    # ─── 6. 커뮤니티 반응 ───
    add_heading_styled(doc, "6. 커뮤니티 반응 시뮬레이션", 1)

    add_body(doc,
        "페르소나들이 체험 후 SNS 스타일로 후기를 공유합니다. "
        "실제 사용자 커뮤니티에서 어떤 반응이 나올지 미리 예측할 수 있습니다. "
        "오른쪽 인사이트 패널에서 전체 감성, NPS, 입소문 예측, 핵심 키워드를 확인합니다.", size=11)

    add_image_full(doc,
        MOCKUPS / "v4-redesign-sns.png",
        "커뮤니티 반응 — 부정 68%, NPS -47. 핵심 키워드: #영어_장벽 14회, #조건부_로직 11회",
        width=Inches(9.5))

    doc.add_page_break()

    # ─── 7. 기술 차별점 ───
    add_heading_styled(doc, "7. 기술 차별점", 1)

    points = [
        ("5-Layer 감정 엔진",
         "OCC → PAD → SDE 수학적 파이프라인. LLM에 감정을 맡기면 sycophancy로 제품에 관대해지는데, "
         "결정론적 계산으로 이를 방지합니다. Ablation 실험으로 검증 완료."),
        ("실제 제품 사용",
         "Playwright로 실제 브라우저를 조작합니다. 설문/Figma/데이터 기반이 아닌, "
         "실제 클릭·입력·네비게이션. 모든 상용 경쟁사가 이 접근을 회피합니다."),
        ("기존 데이터 불필요",
         "Amplitude/Mixpanel 등 기존 분석 데이터 없이, 스테이징 URL만으로 동작합니다. "
         "Pre-launch MVP에 최적화되어 있습니다."),
        ("세그먼트별 분화",
         "같은 제품이어도 anxious/power_user/intuitive 세그먼트별로 "
         "유의미하게 다른 감정·행동 패턴을 보입니다. ANOVA p<0.05 검증 완료."),
    ]

    for title, desc in points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run_t = p.add_run(f"▎ {title}\n")
        run_t.bold = True
        run_t.font.size = Pt(13)
        run_t.font.color.rgb = RGBColor(0xC8, 0x6B, 0x2E)
        run_d = p.add_run(desc)
        run_d.font.size = Pt(11)

    doc.add_page_break()

    # ─── 8. 비즈니스 모델 ───
    add_heading_styled(doc, "8. Go-to-Market", 1)

    add_body(doc, "타겟 고객", bold=True, size=13, color=RGBColor(0xC8, 0x6B, 0x2E))
    add_body(doc,
        "Pre-launch MVP/프로토타입을 가진 스타트업, 프로덕트 팀. "
        "사용자 테스트를 하고 싶지만 시간·비용·리크루팅 장벽이 높은 팀.", size=11)

    add_body(doc, "비즈니스 모델", bold=True, size=13, color=RGBColor(0xC8, 0x6B, 0x2E))
    add_body(doc,
        "• 시뮬레이션 1회 = 페르소나 20명 × 제품 1개 → 리포트 1건\n"
        "• 초기: 컨시어지 모델 (고객 URL + 테스트 계정 → 리포트 딜리버리)\n"
        "• 확장: 셀프서브 SaaS (대시보드에서 직접 실행)", size=11)

    add_body(doc, "현재 단계", bold=True, size=13, color=RGBColor(0xC8, 0x6B, 0x2E))
    add_body(doc,
        "Discovery — 핵심 가설 '시뮬레이션 충실도가 의사결정에 쓸 만한 수준인가' 검증 중. "
        "단일 세션 분화 검증 완료, 다중 세션 시뮬레이션 구현 완료.", size=11)

    # 저장
    doc.save(str(OUTPUT))
    print(f"✅ 저장 완료: {OUTPUT}")
    print(f"   크기: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
